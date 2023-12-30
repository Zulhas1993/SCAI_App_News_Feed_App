import os
import requests
from bs4 import BeautifulSoup
import csv
import docx
from googletrans import Translator
import random

base_url = "https://b.hatena.ne.jp/hotentry/"
categories = ["general", "social", "economics", "it", "living", "learning", "technology", "interesting", "entertainment", "anime and games"]

## Function for Featch html from dedicated website
def fetch_and_save_to_file(url, path):
    try:
        r = requests.get(url)
        if r.status_code == 200:
            os.makedirs(os.path.dirname(path), exist_ok=True)
            with open(path, "w", encoding="utf-8") as f:
                f.write(r.text)
            print("HTML content saved successfully!")
        else:
            print(f"Failed to fetch content. Status code: {r.status_code}")
    except requests.RequestException as e:
        print(f"Request Exception: {e}")
## Function for 
def extract_titles(html_content):
    soup = BeautifulSoup(html_content, 'html.parser')
    titles = soup.find_all('h3', class_='entrylist-contents-title')
    return [title.a.get_text(strip=True) for title in titles]

def translate_to_english(text):
    try:
        translator = Translator()
        translation = translator.translate(text, src='ja', dest='en')
        return translation.text
    except Exception as e:
        print(f"Translation error: {e}")
        return "Translation error"

def save_titles_to_csv(category, titles):
    csv_file_path = f"data/{category}_titles.csv"
    with open(csv_file_path, 'w', newline='', encoding='utf-8') as csvfile:
        csv_writer = csv.writer(csvfile)
        csv_writer.writerow(['Japanese Title', 'English Title'])
        for title in titles:
            english_title = translate_to_english(title)
            csv_writer.writerow([title, english_title])
    print(f"Titles saved to {csv_file_path}")

def generate_theme(keyword):
    title = f"Exploring {keyword}"
    request = f"Can you provide information about {keyword}?"
    return title, request

def save_theme_to_file(title, request, file_path):
    # Ensure the directory exists
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    # Create a new Word document
    doc = docx.Document()

    # Add title and request to the document
    doc.add_heading(title, level=1)
    doc.add_paragraph(request)

    # Save the document to the specified file path
    doc.save(file_path)
    print(f"Theme saved to {file_path}")


def generate_article(theme, scraping_data):
    article_length = 140
    article = f"Exploring {theme}. {random.choice(scraping_data)}"
    return article[:article_length]

def save_article_to_file(article, file_path):
    # Ensure the directory exists
    os.makedirs(os.path.dirname(file_path), exist_ok=True)

    # Save the generated article to a text file
    with open(file_path, 'w', encoding='utf-8') as article_file:
        article_file.write(article)
    print(f"Article saved to {file_path}")


def main():
    user_input = input("Enter a keyword, question, or sentence: ")
    theme_title, theme_request = generate_theme(user_input)

    theme_file_path = "themes/theme_doc.docx"
    article_file_path = "articles/generated_article.txt"

    save_theme_to_file(theme_title, theme_request, theme_file_path)

    scraping_data = ["Data point 1", "Data point 2", "Data point 3", "Data point 4"]
    article = generate_article(theme_title, scraping_data)

    save_article_to_file(article, article_file_path)

    for category in categories:
        url = f"{base_url}{category}"
        html_file_path = f"data/{category}.html"

        fetch_and_save_to_file(url, html_file_path)

        with open(html_file_path, 'r', encoding='utf-8') as html_file:
            html_content = html_file.read()

        titles = extract_titles(html_content)
        print(f"\nTitles for {category.capitalize()} category:")
        for title in titles:
            print(title)

        save_titles_to_csv(category, titles)

if __name__ == "__main__":
    main()
